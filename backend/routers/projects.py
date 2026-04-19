from fastapi import APIRouter, HTTPException, Request
from pydantic import BaseModel
from models.schemas import ProjectCreate, ProjectInfo
from services.session_store import get_project, set_project, get_all_projects, bind_session_to_project
from services.database import (
    get_projects_for_user, add_team_member, remove_team_member,
    get_team_members, get_user_by_username, get_user_by_token,
)
import uuid
import re
from datetime import datetime

router = APIRouter(prefix="/api/projects", tags=["projects"])


def _get_current_user(request: Request) -> dict | None:
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        token = auth[7:]
        return get_user_by_token(token)
    return None


@router.get("/")
def list_projects(request: Request):
    user = _get_current_user(request)
    if user and user.get("role") == "student":
        projects = get_projects_for_user(user["user_id"])
    elif user and user.get("role") == "teacher":
        projects = get_all_projects()
    else:
        projects = get_all_projects()
    return {"projects": projects}


def _build_user_profile(user: dict, force_refresh: bool = False) -> dict:
    """Build the /my-profile response. If force_refresh, regenerate each project's
    capability_profile via LLM; otherwise prefer the cached value on projects row.
    """
    from services.database import (
        get_sessions_for_project,
        get_chat_history as db_get_chat_history,
        get_uploaded_files_for_project,
        save_capability_profile,
    )
    from services.marker_parser import parse_capability_profile
    from config import USE_MOCK_API

    user_projects = get_projects_for_user(user["user_id"])
    if not user_projects:
        return {
            "user": {"username": user["username"], "display_name": user.get("display_name", ""), "class_id": user.get("class_id", "")},
            "projects": [],
            "overall_profile": None,
            "message": "暂无项目，请先创建项目并与AI教练对话",
        }

    project_profiles = []
    all_scores = {"empathy": [], "ideation": [], "business": [], "execution": [], "logic": []}

    for proj in user_projects:
        pid = proj["project_id"]
        sessions = get_sessions_for_project(pid)

        # P3-3: prefer cached profile unless caller forces refresh
        profile = None
        profile_cached = False
        profile_updated_at = proj.get("capability_profile_updated_at")
        if not force_refresh:
            profile = proj.get("capability_profile")
            if profile:
                profile_cached = True

        all_msgs: list[dict] = []
        if not profile_cached:
            for sess in sessions[-3:]:
                all_msgs.extend(db_get_chat_history(sess["session_id"]))

            # Inject uploaded file content so profile LLM sees attachments
            # (file_context is appended to user msg at chat time but not persisted
            # into chat_messages — re-inject from uploaded_files here)
            files = get_uploaded_files_for_project(pid)
            file_texts = [
                f"[文件: {f['filename']}]\n{f['extracted_text'][:2000]}"
                for f in files if f.get("extracted_text")
            ]
            if file_texts and all_msgs:
                file_block = "已上传文件参考资料：\n" + "\n\n".join(file_texts[-3:])
                all_msgs = [{"role": "user", "content": file_block}, *all_msgs]

        profile_from_llm = False
        if not profile and all_msgs:
            if USE_MOCK_API:
                profile = {
                    "empathy": {"score": 7, "evidence": "能描述用户场景", "suggestion": "量化痛点频率"},
                    "ideation": {"score": 6, "evidence": "方案有创意", "suggestion": "拓展差异化"},
                    "business": {"score": 5, "evidence": "模式基本清晰", "suggestion": "补充CAC/LTV"},
                    "execution": {"score": 5, "evidence": "有初步规划", "suggestion": "细化里程碑"},
                    "logic": {"score": 6, "evidence": "论述较清晰", "suggestion": "增加数据支撑"},
                    "overall_comment": "综合表现良好，商业和执行维度有提升空间。",
                }
                profile_from_llm = True
            else:
                try:
                    from services.claude_client import chat_completion
                    _CAP_SYSTEM = """你是一位创新创业教育分析专家。
请根据学生在以下对话记录中的表现，生成一份五维能力画像评分报告。

五个维度（每项 0-10 分，10为最高）：
1. 同理心（Empathy）：识别用户真实痛点、站在用户角度思考的能力
2. 创意发散（Ideation）：提出有创意、有价值的解决方案的能力
3. 商业建模（Business）：理解商业逻辑、收入模型、市场策略的能力
4. 执行规划（Execution）：制定可行计划、分配资源、管理里程碑的能力
5. 逻辑表达（Logic）：清晰论证、数据支撑、结构化表达的能力

先用 1-2 段文字给出整体评价，然后严格按以下 JSON 格式输出（放在代码块中）：

```json
{
  "empathy": {"score": 7, "evidence": "学生提到...", "suggestion": "建议..."},
  "ideation": {"score": 6, "evidence": "...", "suggestion": "..."},
  "business": {"score": 5, "evidence": "...", "suggestion": "..."},
  "execution": {"score": 4, "evidence": "...", "suggestion": "..."},
  "logic": {"score": 6, "evidence": "...", "suggestion": "..."},
  "overall_comment": "综合评价..."
}
```

注意：必须输出完整 JSON，每个维度都要有 score、evidence、suggestion 三个字段。"""
                    # Only send student (user) messages — avoids triggering model
                    # content filters on long/structured assistant replies (rubric JSON etc.)
                    user_msgs = [
                        {"role": "user", "content": (m.get("content") or "")[:800]}
                        for m in all_msgs
                        if m.get("role") == "user" and m.get("content")
                    ][-8:]
                    raw = chat_completion(_CAP_SYSTEM, user_msgs, max_tokens=4096)
                    profile = parse_capability_profile(raw)
                    if profile:
                        profile_from_llm = True
                    if not profile:
                        import logging
                        logging.getLogger(__name__).warning(
                            "capability_profile parse failed for project %s; raw len=%d, tail=%r",
                            pid, len(raw or ""), (raw or "")[-200:]
                        )
                except Exception as _e:
                    import logging
                    logging.getLogger(__name__).warning(
                        "capability_profile LLM call failed for project %s: %s", pid, _e
                    )

        # On force_refresh with LLM failure: keep old cached profile instead of
        # dropping to score-only fallback — preserves trust in "last known good"
        if not profile and force_refresh and proj.get("capability_profile"):
            profile = proj.get("capability_profile")
            profile_cached = True  # displayed data is actually stale cache

        # Fallback: if LLM profile generation failed, use existing project scores
        if not profile and proj.get("scores"):
            s = proj["scores"]
            profile = {
                "empathy": {"score": s.get("empathy", 0), "evidence": "", "suggestion": ""},
                "ideation": {"score": s.get("ideation", 0), "evidence": "", "suggestion": ""},
                "business": {"score": s.get("business", 0), "evidence": "", "suggestion": ""},
                "execution": {"score": s.get("execution", 0), "evidence": "", "suggestion": ""},
                "logic": {"score": s.get("pitching", 0), "evidence": "", "suggestion": ""},
            }

        if profile:
            for dim in all_scores:
                dim_data = profile.get(dim)
                if isinstance(dim_data, dict) and "score" in dim_data:
                    all_scores[dim].append(dim_data["score"])

        # Persist freshly generated LLM profiles only (not fallbacks)
        if profile and not profile_cached and profile_from_llm:
            try:
                save_capability_profile(pid, profile)
                from services.database import get_project as _gp
                _p = _gp(pid)
                if _p:
                    profile_updated_at = _p.get("capability_profile_updated_at")
            except Exception:
                pass

        project_profiles.append({
            "project_id": pid,
            "project_name": proj.get("name", ""),
            "industry": proj.get("industry", ""),
            "stage": proj.get("stage", "discovery"),
            "scores": proj.get("scores"),
            "capability_profile": profile,
            "profile_cached": profile_cached,
            "profile_updated_at": profile_updated_at,
            "session_count": len(sessions),
        })

    # 计算综合画像
    overall = None
    if any(len(v) > 0 for v in all_scores.values()):
        overall = {}
        for dim, vals in all_scores.items():
            overall[dim] = round(sum(vals) / len(vals), 1) if vals else 0

    return {
        "user": {"username": user["username"], "display_name": user.get("display_name", ""), "class_id": user.get("class_id", "")},
        "projects": project_profiles,
        "overall_profile": overall,
    }


@router.get("/my-profile")
def get_my_profile(request: Request):
    """学生端个人画像：默认读缓存，无缓存时即时生成并写回。"""
    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    return _build_user_profile(user, force_refresh=False)


@router.post("/my-profile/refresh")
def refresh_my_profile(request: Request):
    """强制重算当前学生所有项目的能力画像并写回缓存。"""
    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    return _build_user_profile(user, force_refresh=True)


# ── 商业策划书生成 ─────────────────────────────────────────────────

_BP_SYSTEM_PROMPT = """你是一位资深创业顾问，擅长为大学生创业项目撰写完整、专业、可投资的商业策划书。
请严格基于提供的项目数据（对话记录、评分、诊断、采集表单），生成一份结构完整的商业策划书。

# 输出要求
- 输出必须是合法 JSON，字段与下方模板完全一致
- 每个 section 的 content 用 Markdown 格式（可含小标题、列表、表格）
- 内容必须与项目真实情况对齐，不要编造未出现的数据；缺失信息用"待补充"标注
- 针对项目类型（创新项目/商业项目/公益项目）调整盈利模式与社会价值章节的侧重

# 输出 JSON 格式
```json
{
  "title": "项目名 · 商业策划书",
  "project_type": "创新项目 | 商业项目 | 公益项目",
  "executive_summary": "一段不超过 250 字的执行摘要",
  "sections": [
    {"id": "overview", "title": "一、项目概述", "content": "..."},
    {"id": "pain_point", "title": "二、痛点与用户需求", "content": "..."},
    {"id": "solution", "title": "三、解决方案与产品", "content": "..."},
    {"id": "innovation", "title": "四、创新点与技术壁垒", "content": "..."},
    {"id": "market", "title": "五、市场分析与竞品", "content": "..."},
    {"id": "business_model", "title": "六、商业模式与盈利方案", "content": "..."},
    {"id": "marketing", "title": "七、营销与获客策略", "content": "..."},
    {"id": "team", "title": "八、团队介绍", "content": "..."},
    {"id": "financial", "title": "九、财务预测（3年）", "content": "..."},
    {"id": "milestone", "title": "十、执行计划与里程碑", "content": "..."},
    {"id": "risk", "title": "十一、风险分析与应对", "content": "..."},
    {"id": "social_value", "title": "十二、社会价值与影响力", "content": "..."}
  ]
}
```

# 不同项目类型的重点
- **创新项目**：着重"创新点与技术壁垒"、"可行性验证"，商业模式可简洁
- **商业项目**：着重"盈利方案"、"财务预测"、"获客成本(CAC)与客户生命周期价值(LTV)"
- **公益项目**：着重"社会价值"、"受益群体覆盖"、"可持续性（非一次性捐赠）"、"公益与商业平衡"

# 盈利方案建议（根据项目类型选取 2-3 种）
- 订阅制 / SaaS 年费
- 平台佣金 / 交易抽成
- 增值服务 / 会员等级
- 硬件销售 + 软件订阅
- 广告 / 数据合作
- 政府补贴 + 企业定制（公益项目常用）
- 社会影响力债券（SIB）
"""


_VALID_PROJECT_TYPES = {"创新项目", "商业项目", "公益项目"}


def _infer_project_type(project: dict) -> str:
    """根据项目名/行业/描述粗略推断项目类型。"""
    text = (project.get("name", "") + " " + project.get("industry", "") + " " + project.get("description", "")).lower()
    if any(k in text for k in ["公益", "乡村", "扶贫", "助残", "养老", "社会创新", "社会企业"]):
        return "公益项目"
    if any(k in text for k in ["创新", "技术", "研发", "专利", "实验室", "算法"]):
        return "创新项目"
    return "商业项目"


# Project-type-specific guidance for coach prompt and BP generation
PROJECT_TYPE_FOCUS = {
    "创新项目": (
        "【项目类型：创新项目】重点关注：\n"
        "- 创新点与技术壁垒（专利、算法、工艺）\n"
        "- 可行性验证（MVP、实验数据、概念验证）\n"
        "- 研发路线图与团队技术背景\n"
        "- 盈利模式可简洁（专利授权、技术转让、合作研发）"
    ),
    "商业项目": (
        "【项目类型：商业项目】重点关注：\n"
        "- 明确的盈利模式（订阅、抽成、广告、增值服务）\n"
        "- 获客成本(CAC) 与 客户生命周期价值(LTV)\n"
        "- 市场规模、竞品分析、差异化定位\n"
        "- 3年财务预测与融资计划"
    ),
    "公益项目": (
        "【项目类型：公益项目】重点关注：\n"
        "- 社会价值与受益群体覆盖面\n"
        "- 可持续性（非一次性捐赠，需有造血机制）\n"
        "- 公益与商业的平衡（社会企业、SIB 社会影响力债券）\n"
        "- 主要收入来源：政府采购、基金会资助、企业 CSR 合作"
    ),
}


# ── BP Readiness Gate ──────────────────────────────────────────────
# 标准门槛：
#   1. score_snapshots ≥ 3（至少 3 轮有评分的对话）
#   2. 累计用户消息数 ≥ 6
#   3. 5 维评分至少 3 维 > 0
MIN_SCORE_ROUNDS = 3
MIN_USER_MESSAGES = 6
MIN_SCORED_DIMS = 3


def _check_bp_readiness(project_id: str) -> dict:
    """检查项目是否满足生成策划书的最低证据门槛。"""
    from services.database import (
        get_score_snapshots as _gss,
        get_sessions_for_project as _gsp,
        get_chat_history as _gh,
    )
    snapshots = _gss(project_id)
    round_count = len(snapshots)

    # 累计用户消息数
    user_msg_count = 0
    for sess in _gsp(project_id):
        msgs = _gh(sess["session_id"])
        user_msg_count += sum(1 for m in msgs if m.get("role") == "user")

    # 5 维评分覆盖
    proj = get_project(project_id)
    scores = (proj or {}).get("scores") or {}
    scored_dims = sum(1 for v in scores.values() if isinstance(v, (int, float)) and v > 0)

    ok = (
        round_count >= MIN_SCORE_ROUNDS
        and user_msg_count >= MIN_USER_MESSAGES
        and scored_dims >= MIN_SCORED_DIMS
    )
    missing = []
    if round_count < MIN_SCORE_ROUNDS:
        missing.append(f"有效评分轮数 {round_count}/{MIN_SCORE_ROUNDS}")
    if user_msg_count < MIN_USER_MESSAGES:
        missing.append(f"对话消息数 {user_msg_count}/{MIN_USER_MESSAGES}")
    if scored_dims < MIN_SCORED_DIMS:
        missing.append(f"已评估维度 {scored_dims}/{MIN_SCORED_DIMS}")

    return {
        "ready": ok,
        "round_count": round_count,
        "user_msg_count": user_msg_count,
        "scored_dims": scored_dims,
        "requirements": {
            "min_rounds": MIN_SCORE_ROUNDS,
            "min_messages": MIN_USER_MESSAGES,
            "min_dims": MIN_SCORED_DIMS,
        },
        "missing": missing,
    }


@router.get("/{project_id}/bp/readiness")
def get_bp_readiness(project_id: str, request: Request):
    """返回策划书生成门槛的进度。"""
    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if user.get("role") == "student" and project.get("owner_id") != user["user_id"]:
        raise HTTPException(status_code=403, detail="无权访问")
    return _check_bp_readiness(project_id)


@router.post("/{project_id}/generate-bp")
def generate_business_plan(project_id: str, request: Request):
    """生成商业策划书（基于项目所有对话、评分、诊断）。"""
    from services.database import (
        get_sessions_for_project,
        get_chat_history as _gh,
        save_project,
        get_uploaded_files_for_project as _gf,
    )
    from services.claude_client import chat_completion
    from config import USE_MOCK_API
    import json as _json

    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")

    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")

    # 权限校验：owner 或 admin/teacher
    if user.get("role") == "student" and project.get("owner_id") != user["user_id"]:
        raise HTTPException(status_code=403, detail="无权生成他人项目的策划书")

    # 证据门槛校验（学生身份强制，教师/管理员可绕过用于演示）
    if user.get("role") == "student":
        readiness = _check_bp_readiness(project_id)
        if not readiness["ready"]:
            raise HTTPException(
                status_code=400,
                detail="证据不足，无法生成策划书：" + "；".join(readiness["missing"]) +
                       "。请继续与AI教练对话，积累证据后再生成。",
            )

    # Use explicit project_type if set, else auto-infer
    project_type = project.get("project_type") if project.get("project_type") in _VALID_PROJECT_TYPES else _infer_project_type(project)

    # 聚合对话记录
    sessions = get_sessions_for_project(project_id)
    all_msgs: list[dict] = []
    for sess in sessions[-5:]:
        all_msgs.extend(_gh(sess["session_id"]))

    scores = project.get("scores", {})
    diagnosis = project.get("diagnosis", [])
    rubric_full = project.get("rubric_full", {})

    # 构造输入上下文
    context_lines = [
        f"项目名称：{project.get('name', '')}",
        f"行业：{project.get('industry', '')}",
        f"项目描述：{project.get('description', '')}",
        f"当前阶段：{project.get('stage', 'discovery')}",
        f"推断项目类型：{project_type}",
    ]
    if scores:
        context_lines.append(f"五维评分：{_json.dumps(scores, ensure_ascii=False)}")
    if diagnosis:
        context_lines.append(f"诊断问题：{'; '.join(diagnosis[:5])}")
    if rubric_full:
        context_lines.append(f"Rubric 详评（简要）：{_json.dumps({k: v.get('score') for k, v in rubric_full.items() if isinstance(v, dict)}, ensure_ascii=False)}")

    type_focus = PROJECT_TYPE_FOCUS.get(project_type, "")

    # 注入学生上传的文件正文（调研报告 / BP 草稿等），避免信息丢失
    _files = _gf(project_id)
    _file_texts = [
        f"[文件: {f['filename']}]\n{f['extracted_text'][:2000]}"
        for f in _files if f.get("extracted_text")
    ]
    file_section = (
        "\n\n【已上传文件参考资料（最近3份）】\n" + "\n\n".join(_file_texts[-3:])
        if _file_texts else ""
    )

    user_msg = {
        "role": "user",
        "content": "【项目元数据】\n" + "\n".join(context_lines) +
                   (f"\n\n{type_focus}" if type_focus else "") +
                   file_section +
                   "\n\n【最近对话摘录（最多5个会话）】\n" +
                   "\n".join([f"{m.get('role','')}: {(m.get('content','') or '')[:300]}" for m in all_msgs[-30:]]) +
                   "\n\n请为此项目生成完整商业策划书，严格按系统提示的 JSON 格式输出。",
    }

    # Mock / Fallback 生成
    def _fallback_bp() -> dict:
        return {
            "title": f"{project.get('name','项目')} · 商业策划书",
            "project_type": project_type,
            "executive_summary": project.get("description", "") or "执行摘要待完善。",
            "sections": [
                {"id": "overview", "title": "一、项目概述", "content": f"**行业**：{project.get('industry','')}\n\n**简介**：{project.get('description','')}"},
                {"id": "pain_point", "title": "二、痛点与用户需求", "content": "待补充（请先与AI教练深入对话）"},
                {"id": "solution", "title": "三、解决方案与产品", "content": "待补充"},
                {"id": "innovation", "title": "四、创新点与技术壁垒", "content": "待补充"},
                {"id": "market", "title": "五、市场分析与竞品", "content": "待补充"},
                {"id": "business_model", "title": "六、商业模式与盈利方案", "content": "待补充"},
                {"id": "marketing", "title": "七、营销与获客策略", "content": "待补充"},
                {"id": "team", "title": "八、团队介绍", "content": "待补充"},
                {"id": "financial", "title": "九、财务预测（3年）", "content": "待补充"},
                {"id": "milestone", "title": "十、执行计划与里程碑", "content": "待补充"},
                {"id": "risk", "title": "十一、风险分析与应对", "content": f"**已识别诊断问题**：{'; '.join(diagnosis) if diagnosis else '暂无'}"},
                {"id": "social_value", "title": "十二、社会价值与影响力", "content": "待补充"},
            ],
        }

    bp: dict | None = None
    if USE_MOCK_API:
        bp = _fallback_bp()
    else:
        try:
            raw = chat_completion(_BP_SYSTEM_PROMPT, [user_msg], max_tokens=6000)
            # 提取 JSON（兼容 ```json ... ``` 包裹）
            import re as _re
            m = _re.search(r"\{[\s\S]*\}", raw)
            if m:
                bp = _json.loads(m.group(0))
        except Exception:
            bp = None

    if not bp:
        bp = _fallback_bp()

    bp["generated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # 持久化到项目
    project["bp_content"] = bp
    save_project(project)

    return {"ok": True, "bp": bp}


@router.get("/{project_id}/bp")
def get_business_plan(project_id: str, request: Request):
    """读取已生成的商业策划书。"""
    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if user.get("role") == "student" and project.get("owner_id") != user["user_id"]:
        raise HTTPException(status_code=403, detail="无权访问")
    bp = project.get("bp_content")
    if not bp:
        return {"exists": False}
    return {"exists": True, "bp": bp}


def _add_markdown_para(doc, text: str, style: str = None):
    import re
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        if i % 2 == 1:
            run.bold = True


@router.get("/{project_id}/bp/download")
def download_business_plan(project_id: str, request: Request):
    """下载商业策划书为 docx 文件。"""
    from fastapi.responses import StreamingResponse
    import io

    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if user.get("role") == "student" and project.get("owner_id") != user["user_id"]:
        raise HTTPException(status_code=403, detail="无权下载")
    bp = project.get("bp_content")
    if not bp:
        raise HTTPException(status_code=404, detail="尚未生成策划书，请先点击生成")

    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        raise HTTPException(status_code=500, detail="服务端缺少 python-docx，请运行 pip install python-docx")

    doc = Document()
    title = doc.add_heading(bp.get("title", "商业策划书"), level=0)
    doc.add_paragraph(f"项目类型：{bp.get('project_type', '')}")
    doc.add_paragraph(f"生成时间：{bp.get('generated_at', '')}")
    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph(bp.get("executive_summary", ""))

    for section in bp.get("sections", []):
        doc.add_heading(section.get("title", ""), level=1)
        content = section.get("content", "")
        for para in content.split("\n"):
            para = para.strip()
            if not para:
                continue
            # Markdown headings → Word heading styles
            if para.startswith("#### "):
                doc.add_heading(para[5:], level=4)
            elif para.startswith("### "):
                doc.add_heading(para[4:], level=3)
            elif para.startswith("## "):
                doc.add_heading(para[3:], level=2)
            elif para.startswith("# "):
                doc.add_heading(para[2:], level=1)
            elif para.startswith("- ") or para.startswith("* "):
                # Bullet: handle inline **bold**
                _add_markdown_para(doc, para[2:], style="List Bullet")
            else:
                _add_markdown_para(doc, para)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = (project.get("name", "project") or "project").replace("/", "_").replace("\\", "_")
    filename = f"{safe_name}_BP.docx"
    # URL-encode filename for non-ASCII chars
    from urllib.parse import quote
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(filename)}"},
    )


@router.get("/{project_id}")
def get_project_detail(project_id: str):
    proj = get_project(project_id)
    if not proj:
        raise HTTPException(status_code=404, detail="项目不存在")
    return proj


@router.post("/", response_model=ProjectInfo)
def create_project(req: ProjectCreate, request: Request, owner_id: str = "student_001"):
    user = _get_current_user(request)
    if user:
        owner_id = user["user_id"]
    project_id = f"proj_{uuid.uuid4().hex[:6]}"
    # Resolve project_type: use explicit value if valid, else auto-infer
    ptype = req.project_type if req.project_type in _VALID_PROJECT_TYPES else None
    project = {
        "project_id": project_id,
        "name": req.name,
        "industry": req.industry,
        "description": req.description,
        "stage": "discovery",
        "owner_id": owner_id,
        "scores": {},
        "diagnosis": [],
        "project_type": ptype or _infer_project_type({"name": req.name, "industry": req.industry, "description": req.description}),
        "created_at": datetime.now().strftime("%Y-%m-%d"),
    }
    set_project(project_id, project)
    # Auto-add owner as team member
    add_team_member(project_id, owner_id, "owner")
    return ProjectInfo(**project)


class ProjectTypeUpdate(BaseModel):
    project_type: str


@router.patch("/{project_id}/type")
def update_project_type(project_id: str, req: ProjectTypeUpdate, request: Request):
    """允许学生本人（或教师/管理员）手动修正项目类型。"""
    from services.database import save_project
    user = _get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="请先登录")
    if req.project_type not in _VALID_PROJECT_TYPES:
        raise HTTPException(status_code=400, detail=f"项目类型必须为 {list(_VALID_PROJECT_TYPES)} 之一")
    project = get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if user.get("role") == "student" and project.get("owner_id") != user["user_id"]:
        raise HTTPException(status_code=403, detail="无权修改他人项目")
    project["project_type"] = req.project_type
    save_project(project)
    return {"ok": True, "project_type": req.project_type}


# ── Auto-infer project from chat text ──────────────────────────────

class AutoInferRequest(BaseModel):
    text: str


# Industry keywords for heuristic matching
_INDUSTRY_PATTERNS = [
    (["医疗", "健康", "诊断", "眼科", "病", "患者", "医院", "药"], "医疗健康"),
    (["教育", "学习", "学生", "课程", "培训", "考试"], "教育科技"),
    (["农业", "农村", "种植", "养殖", "农民"], "农业科技"),
    (["金融", "理财", "投资", "银行", "保险", "支付"], "金融科技"),
    (["物流", "仓储", "配送", "运输", "供应链"], "物流供应链"),
    (["环保", "碳", "新能源", "光伏", "储能", "绿色"], "节能环保"),
    (["零售", "电商", "购物", "消费", "商城"], "新零售"),
    (["餐饮", "食品", "外卖", "烹饪", "厨房"], "餐饮食品"),
    (["房地产", "建筑", "装修", "物业"], "房地产建筑"),
    (["旅游", "酒店", "民宿", "景区", "出行"], "文旅出行"),
    (["游戏", "娱乐", "直播", "内容", "短视频"], "文娱传媒"),
    (["制造", "工业", "自动化", "机器人", "生产"], "先进制造"),
    (["AI", "人工智能", "大模型", "机器学习", "深度学习"], "人工智能"),
]


def _heuristic_infer(text: str) -> dict:
    """Extract project name and industry from free text using heuristics."""
    # Try to find project name patterns
    name_patterns = [
        r"(做|做一个|开发|研发|想做|计划做)\s*[一个]?\s*([^\s，。！？、,!?]{4,20}(?:系统|平台|APP|应用|产品|工具|服务|方案))",
        r"([^\s，。！？、,!?]{2,15}(?:系统|平台|APP|应用|产品|工具|服务|方案))",
        r"(我的项目[是叫]?|项目名[称叫]?|项目[是叫]?)[：:「]?\s*([^\s，。！？「」,!?]{3,20})",
    ]
    name = ""
    for pattern in name_patterns:
        m = re.search(pattern, text)
        if m:
            name = m.group(2) if len(m.groups()) >= 2 else m.group(1)
            name = name.strip("「」""''")
            if len(name) >= 4:
                break

    # Fallback: take first sentence as rough name
    if not name:
        first = re.split(r"[，。！？\n]", text)[0].strip()
        if 5 <= len(first) <= 25:
            name = first

    # Detect industry
    industry = ""
    for keywords, label in _INDUSTRY_PATTERNS:
        if any(kw in text for kw in keywords):
            industry = label
            break

    return {"name": name[:30] if name else "", "industry": industry, "description": ""}


@router.post("/auto-infer")
def auto_infer_project(req: AutoInferRequest, request: Request):
    """
    Infer project name and industry from free text.
    First tries LLM, falls back to heuristics.
    """
    from config import USE_MOCK_API
    result = {"name": "", "industry": "", "description": ""}

    if not USE_MOCK_API:
        try:
            from services.claude_client import chat_completion
            from config import MODEL_LIGHT
            system = """从用户描述中提取创业项目信息。
只输出 JSON，格式：{"name": "项目名称", "industry": "行业", "description": "一句话描述"}
- name: 4-20字，描述项目的核心产品/服务，不要包含"我的项目"等词
- industry: 从以下选择：医疗健康/教育科技/农业科技/金融科技/物流供应链/节能环保/新零售/餐饮食品/人工智能/先进制造/文旅出行/文娱传媒/其他
- description: 一句话说明项目价值主张
如果信息不足，name 输出空字符串""。"""
            msgs = [{"role": "user", "content": req.text}]
            raw = chat_completion(system, msgs, model=MODEL_LIGHT)
            import json
            m = re.search(r'\{.*?\}', raw, re.DOTALL)
            if m:
                parsed = json.loads(m.group(0))
                result = {
                    "name": parsed.get("name", "")[:30],
                    "industry": parsed.get("industry", ""),
                    "description": parsed.get("description", ""),
                }
        except Exception:
            pass

    # Fallback to heuristics if LLM failed or returned empty name
    if not result.get("name"):
        result = _heuristic_infer(req.text)

    return result


# ── Bind session to project ─────────────────────────────────────────

class BindSessionRequest(BaseModel):
    session_id: str
    project_id: str


@router.post("/bind-session")
def bind_session(req: BindSessionRequest):
    bind_session_to_project(req.session_id, req.project_id)
    return {"ok": True}


# ── Team Management (F6-adv) ────────────────────────────────────────

class AddMemberRequest(BaseModel):
    username: str


@router.get("/{project_id}/team")
def get_team(project_id: str):
    members = get_team_members(project_id)
    return {"project_id": project_id, "members": members}


@router.post("/{project_id}/team")
def add_member(project_id: str, req: AddMemberRequest):
    user = get_user_by_username(req.username)
    if not user:
        raise HTTPException(status_code=404, detail=f"用户 '{req.username}' 不存在")
    add_team_member(project_id, user["user_id"], "member")
    return {"ok": True, "user_id": user["user_id"], "username": req.username}


@router.delete("/{project_id}/team/{user_id}")
def remove_member(project_id: str, user_id: str):
    remove_team_member(project_id, user_id)
    return {"ok": True}


# ── 学生个人画像 ─────────────────────────────────────────────────────

